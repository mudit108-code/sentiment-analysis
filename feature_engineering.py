import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, classification_report
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document

# Function for Text Preprocessing
def preprocess_text(text):
    # Your text preprocessing code here
    return text

# Function for Feature Engineering using TF-IDF
def perform_feature_engineering(df):
    # Split the dataset into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(df['processed_review'], df['sentiment'], test_size=0.2, random_state=42)

    # TF-IDF Vectorization
    vectorizer = TfidfVectorizer(max_features=5000)
    X_train_tfidf = vectorizer.fit_transform(X_train)
    X_test_tfidf = vectorizer.transform(X_test)

    return X_train_tfidf, X_test_tfidf, y_train, y_test

# Load the preprocessed dataset
df = pd.read_csv("preprocessed_imdb_dataset.csv")

# Apply text preprocessing to the 'review' column
df['processed_review'] = df['review'].apply(preprocess_text)

# Perform feature engineering
X_train_tfidf, X_test_tfidf, y_train, y_test = perform_feature_engineering(df)

# Initialize and train a logistic regression model
model = LogisticRegression()
model.fit(X_train_tfidf, y_train)

# Make predictions on the test set
y_pred = model.predict(X_test_tfidf)

# Evaluate the model
accuracy = accuracy_score(y_test, y_pred)
print(f"Accuracy: {accuracy}")

# Get the classification report as a string
classification_report_str = classification_report(y_test, y_pred)

# Save the classification report to a Word file
document = Document()
document.add_heading('Classification Report', 0)
document.add_paragraph(classification_report_str)
document.save('classification_report.docx')

# Continue with the rest of your code...
