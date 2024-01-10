import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import classification_report
from docx import Document

# Function for Text Preprocessing
def preprocess_text(text):
    # Your text preprocessing code here
    return text

# Function for Feature Engineering using TF-IDF
def perform_feature_engineering(df):
    # Split the dataset into training and testing sets
    X_train, X_test, y_train, y_test = train_test_split(df['processed_review'], df['sentiment'], test_size=0.2, random_state=42)

    # TF-IDF Vectorization
    vectorizer = TfidfVectorizer(max_features=5000)
    X_train_tfidf = vectorizer.fit_transform(X_train)
    X_test_tfidf = vectorizer.transform(X_test)

    return X_train_tfidf, X_test_tfidf, y_train, y_test

# Load the preprocessed dataset
df = pd.read_csv("preprocessed_imdb_dataset.csv")

# Apply text preprocessing to the 'review' column
df['processed_review'] = df['review'].apply(preprocess_text)

# Perform feature engineering
X_train_tfidf, X_test_tfidf, y_train, y_test = perform_feature_engineering(df)

# Hyperparameter Tuning using Grid Search
param_grid = {
    'C': [0.001, 0.01, 0.1, 1, 10, 100],
    'penalty': ['l1', 'l2']
}

# Initialize logistic regression model
model = LogisticRegression()

# Grid Search with Cross-Validation
grid_search = GridSearchCV(model, param_grid, cv=5, scoring='accuracy')
grid_search.fit(X_train_tfidf, y_train)

# Get the best hyperparameters
best_hyperparameters = grid_search.best_params_

# Evaluate the model with the best hyperparameters on the test set
best_model = grid_search.best_estimator_
y_pred = best_model.predict(X_test_tfidf)

# Get the classification report as a string
classification_report_str = classification_report(y_test, y_pred)

# Save the report to a Word file
document = Document()
document.add_heading('Final Report', level=1)

# Add best hyperparameters to the document
document.add_heading('Best Hyperparameters:', level=2)
for key, value in best_hyperparameters.items():
    document.add_paragraph(f'{key}: {value}')

# Add classification report to the document
document.add_heading('Classification Report:', level=2)
document.add_paragraph(classification_report_str)

# Save the Word document
document.save('final_report.docx')
